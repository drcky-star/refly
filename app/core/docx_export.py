"""Kaynakça Word (.docx) çıktısı — EndNote'un 'bibliography' işlevi."""
from __future__ import annotations
import io
from docx import Document
from docx.shared import Pt, Inches

from . import csl


def _hanging(paragraph):
    """Akademik kaynakça için asılı girinti (ilk satır solda, devamı içeride)."""
    pf = paragraph.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.first_line_indent = Inches(-0.5)
    pf.space_after = Pt(6)


def bibliography_docx(records: list[dict], style: str = "vancouver",
                      title: str = "Kaynaklar", hanging: bool = True) -> bytes:
    """Seçilen stilde numaralı, asılı girintili kaynakça içeren .docx üretir."""
    entries = csl.build_reference_list(records, style=style)
    doc = Document()
    style_n = doc.styles["Normal"]
    style_n.font.name = "Calibri"
    style_n.font.size = Pt(11)

    if title:
        doc.add_heading(title, level=1)

    for line in entries:
        p = doc.add_paragraph(line)
        if hanging:
            _hanging(p)
        else:
            p.paragraph_format.space_after = Pt(6)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
