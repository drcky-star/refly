"""Atıf aracı (cite-while-write lite).

Kullanıcı metnine {{id}}, {{cite:id}} veya [#id] biçiminde işaretçi koyar (id =
kütüphanedeki referans no). Bu modül:
  1) işaretçileri ilk-görülme sırasına göre [1], [2], ... numaralar,
  2) yalnızca atıf yapılan kaynakları o sırayla toplar,
  3) seçilen stilde numaralı kaynakça üretir,
  4) gövde + kaynakça içeren bir Word (.docx) döndürür.
"""
from __future__ import annotations
import io
import re

from docx import Document
from docx.shared import Pt

from . import csl

_MARKER = re.compile(r"\{\{\s*(?:cite:)?(\d+)\s*\}\}|\[#(\d+)\]")


def _used_ids_in_order(text: str) -> list[int]:
    order: list[int] = []
    for m in _MARKER.finditer(text):
        rid = int(m.group(1) or m.group(2))
        if rid not in order:
            order.append(rid)
    return order


def renumber(text: str, id_to_num: dict[int, int]) -> str:
    def repl(m):
        rid = int(m.group(1) or m.group(2))
        n = id_to_num.get(rid)
        return f"[{n}]" if n else ""
    return _MARKER.sub(repl, text)


def process(text: str, get_ref, style: str = "vancouver") -> dict:
    """get_ref: id -> kayıt sözlüğü (yoksa None). Döner: {text, entries, missing}."""
    order = _used_ids_in_order(text)
    records, kept_ids, missing = [], [], []
    for rid in order:
        r = get_ref(rid)
        if r:
            records.append(r)
            kept_ids.append(rid)
        else:
            missing.append(rid)
    id_to_num = {rid: i + 1 for i, rid in enumerate(kept_ids)}
    new_text = renumber(text, id_to_num)
    entries = csl.build_reference_list(records, style=style)
    return {"text": new_text, "entries": entries, "missing": missing,
            "n_citations": len(kept_ids)}


def to_docx(result: dict, title: str = "") -> bytes:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    if title:
        doc.add_heading(title, level=0)
    for para in result["text"].split("\n"):
        p = doc.add_paragraph(para)
        p.paragraph_format.space_after = Pt(6)
    if result["entries"]:
        doc.add_heading("Kaynaklar", level=1)
        for line in result["entries"]:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
